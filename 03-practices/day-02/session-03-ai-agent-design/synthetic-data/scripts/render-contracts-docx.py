#!/usr/bin/env python3
"""
Tạo 4 file DOCX hợp đồng mô phỏng theo Nghị định 30/2020/NĐ-CP.
Font: Times New Roman 14pt, margins 2-3-2-2, line spacing 1.5, justify alignment.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def setup_document():
    """Cấu hình document theo Nghị định 30."""
    doc = Document()

    # Margins: Top 2cm, Left 3cm, Bottom 2cm, Right 2cm
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Set Vietnamese font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = doc.styles['Normal'].element.get_or_add_rPr()
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

    return doc


def add_title(doc, text):
    """Tiêu đề chính: căn giữa, bold, in hoa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_contract_number(doc, text):
    """Số hợp đồng: căn giữa, in nghiêng."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    run.italic = True


def add_disclaimer(doc):
    """Dòng mô phỏng: căn giữa, in nghiêng, xám."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("[Mô phỏng — không sử dụng dữ liệu thật]")
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_party_info(doc, label, info_lines):
    """Thông tin bên A/B: căn trái, bold cho label."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    for line in info_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'


def add_article_title(doc, text):
    """Tiêu đề điều khoản: bold, căn trái."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'


def add_paragraph_text(doc, text, indent=0, italic=False):
    """Đoạn văn bản thường: căn đều."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    if italic:
        run.italic = True
    return p


def add_bullet(doc, text, indent=1):
    """Bullet point."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("- " + text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'


def add_signature_block(doc, date_text):
    """Khối ký kết."""
    # Dòng trống
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date_text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Hai cột chữ ký
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐẠI DIỆN BÊN A")  # Không gạch chân
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.add_run("                    ")
    run2 = p.add_run("ĐẠI DIỆN BÊN B")  # Không gạch chân
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(Ký, đóng dấu)")
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.add_run("            ")
    run2 = p.add_run("(Ký, đóng dấu)")
    run2.font.size = Pt(13)
    run2.font.name = 'Times New Roman'
    run2.italic = True


def build_contract_001():
    """Hợp đồng 001: truyền dẫn liên quận, bình thường."""
    doc = setup_document()

    add_title(doc, "HỢP ĐỒNG CUNG CẤP DỊCH VỤ\nTRUYỀN DẪN LIÊN QUẬN")
    add_contract_number(doc, "Số: HD-DV-2026-001")
    add_disclaimer(doc)

    add_party_info(doc, "Bên A:", [
        "Công ty Cổ phần Viễn thông Mô phỏng (VTN-Sim)",
        "Địa chỉ: 123 đường Nguyễn Văn Cừ, Quận 5, TP. HCM",
        "Mã số thuế: 0301234567-001 (mô phỏng)",
        "Người đại diện: Nguyễn Văn An — Giám đốc",
    ])
    add_paragraph_text(doc, "")
    add_party_info(doc, "Bên B:", [
        "Công ty TNHH Giải pháp Số Demo",
        "Địa chỉ: 456 đường Lê Lợi, Quận 1, TP. HCM",
        "Mã số thuế: 0309876543-001 (mô phỏng)",
        "Người đại diện: Trần Thị Bình — Trưởng phòng Mua hàng",
    ])

    add_paragraph_text(doc, "")

    add_article_title(doc, "ĐIỀU 1: ĐỐI TƯỢNG HỢP ĐỒNG")
    add_paragraph_text(doc, "Bên A cung cấp cho Bên B dịch vụ truyền dẫn liên quận với băng thông 100 Mbps, bao gồm:")
    add_bullet(doc, "Truyền dẫn dữ liệu giữa các điểm kết nối tại Quận 1 và Quận 5")
    add_bullet(doc, "Cam kết sẵn sàng 99.5% thời gian hoạt động")
    add_bullet(doc, "Hỗ trợ kỹ thuật 24/7")

    add_article_title(doc, "ĐIỀU 2: THỜI HẠN HỢP ĐỒNG")
    add_paragraph_text(doc, "Hợp đồng có hiệu lực kể từ ngày 01 tháng 01 năm 2026.")
    add_paragraph_text(doc, "Hợp đồng có thời hạn 12 (mười hai) tháng, hết hạn vào ngày 31 tháng 12 năm 2026.")
    add_paragraph_text(doc, "Không tự động gia hạn. Các bên cần thỏa thuận gia hạn bằng văn bản trước 30 ngày.")

    add_article_title(doc, "ĐIỀU 3: GIÁ TRỊ VÀ THANH TOÁN")
    add_paragraph_text(doc, "Giá trị hợp đồng: 360.000.000 VNĐ (Ba trăm sáu mươi triệu đồng).")
    add_paragraph_text(doc, "Phương thức thanh toán: hàng quý, trong vòng 15 ngày kể từ ngày nhận hóa đơn.")
    add_paragraph_text(doc, "Mỗi quý thanh toán: 90.000.000 VNĐ.")

    add_article_title(doc, "ĐIỀU 4: ĐIỀU KHOẢN PHẠT VI PHẠM")
    add_paragraph_text(doc, "Nếu Bên A không đảm bảo cam kết sẵn sàng 99.5% trong một tháng:")
    add_bullet(doc, "Phạt 1% giá trị hợp đồng hàng tháng cho mỗi 0.5% giảm hiệu suất.")
    add_bullet(doc, "Tối đa phạt không quá 10% giá trị hợp đồng hàng quý.")
    add_paragraph_text(doc, "Nếu Bên B thanh toán chậm:")
    add_bullet(doc, "Phạt 0.05% giá trị chưa thanh toán cho mỗi ngày chậm.")
    add_bullet(doc, "Sau 30 ngày chậm, Bên A có quyền tạm ngưng dịch vụ sau khi thông báo trước 5 ngày.")

    add_article_title(doc, "ĐIỀU 5: BẢO MẬT THÔNG TIN")
    add_paragraph_text(doc, "Cả hai bên cam kết bảo mật mọi thông tin trao đổi trong quá trình thực hiện hợp đồng.")
    add_paragraph_text(doc, "Thông tin bảo mật bao gồm: dữ liệu kỹ thuật, thông tin khách hàng, giá cả và điều khoản thương mại.")
    add_paragraph_text(doc, "Nghĩa vụ bảo mật kéo dài 2 năm sau khi hợp đồng hết hạn.")

    add_article_title(doc, "ĐIỀU 6: CHẤM DỨT HỢP ĐỒNG")
    add_paragraph_text(doc, "Hợp đồng chấm dứt trong các trường hợp:")
    add_bullet(doc, "Hết hạn theo Điều 2 mà không gia hạn.")
    add_bullet(doc, "Thỏa thuận bằng văn bản của cả hai bên.")
    add_bullet(doc, "Một bên vi phạm nghiêm trọng và không khắc phục trong 30 ngày sau khi nhận thông báo.")

    add_article_title(doc, "ĐIỀU 7: GIẢI QUYẾT TRANH CHẤP")
    add_paragraph_text(doc, "Mọi tranh chấp phát sinh sẽ được giải quyết qua thương lượng.")
    add_paragraph_text(doc, "Nếu không thương lượng được trong 30 ngày, tranh chấp sẽ được đưa ra Trọng tài thương mại TP. HCM.")

    add_signature_block(doc, "Ngày ký: 15 tháng 12 năm 2025")

    return doc


def build_contract_002():
    """Hợp đồng 002: mua thiết bị, lỗi OCR, thiếu trường."""
    doc = setup_document()

    add_title(doc, "HỢP ĐỒNG MUA SẮP THIẾT BỊ MẠNG")
    add_contract_number(doc, "Số: HD-MS-2026-002")
    add_disclaimer(doc)

    add_party_info(doc, "Bên A:", [
        "Công ty Cổ phần Viễn thông Mô phỏng (VTN-Sim)",
        "Địa chỉ: 123 đường Nguyễn Văn Cừ, Quận 5, TP. HCM",
        "Mã số thuế: 0301234567-001 (mô phỏng)",
        "Người đại diện: Lê Văn Châu — Phó giám đốc",
    ])
    add_paragraph_text(doc, "")
    add_party_info(doc, "Bên B:", [
        "Công ty Cổ phần Công nghệ Demo",
        "Địa chỉ: 789 đường Trần Hưng Đạo, Quận 5, TP. HCM",
        "Mã số thuế: 0305678901-001 (mô phỏng)",
        "Người đại diện: Phạm Thị Dung — Giám đốc",
    ])

    add_paragraph_text(doc, "")

    add_article_title(doc, "ĐIỀU 1: ĐỐI TƯỢNG HỢP ĐỒNG")
    add_paragraph_text(doc, "Bên B cung cấp cho Bên A:")
    add_bullet(doc, "10 bộ chuyển mạch (switch) Layer 3, model XS-4800 (mô phỏng)")
    add_bullet(doc, "5 bộ định tuyến (router) cỡ aggregation, model XR-3600 (mô phỏng)")
    add_bullet(doc, "Bảo hành 36 tháng")

    add_article_title(doc, "ĐIỀU 2: THỜI HẠN HỢP ĐỒNG")
    add_paragraph_text(doc, "Hợp đồng có hiệu lực kể từ ngày 15 tháng 02 năm 2026.")
    add_paragraph_text(doc, "[ĐIỀU KHOẢN VỀ NGÀY HẾT HẠN BỊ THIẾU DO LỖI OCR]", italic=True)

    add_article_title(doc, "ĐIỀU 3: GIÁ TRỊ VÀ THANH TOÁN")
    add_paragraph_text(doc, "Giá trị hợp đồng: 2.500.000.000 VNĐ (Hai tỷ năm trăm triệu đồng).")
    add_paragraph_text(doc, "Thanh toán theo giai đoạn:")
    add_bullet(doc, "30% khi ký hợp đồng: 750.000.000 VNĐ")
    add_bullet(doc, "40% khi giao hàng: 1.000.000.000 VNĐ")
    add_bullet(doc, "30% sau nghiệm thu: 750.000.000 VNĐ")

    add_article_title(doc, "ĐIỀU 4: ĐIỀU KHOẢN PHẠT")
    add_paragraph_text(doc, "Nếu Bên B giao hàng chậm:")
    add_bullet(doc, "Phạt 0.1% giá trị lô hàng chậm cho mỗi ngày chậm.")
    add_bullet(doc, "Tối đa phạt không quá 5% giá trị hợp đồng.")
    add_paragraph_text(doc, "Nếu Bên A nghiệm thu chậm:")
    add_bullet(doc, "Phạt 0.03% giá trị lô hàng cho mỗi ngày chậm sau 15 ngày kể từ ngày giao hàng.")

    add_article_title(doc, "ĐIỀU 5: BẢO MẬT")
    add_paragraph_text(doc, "[ĐIỀU KHOẢN BẢO MẬT BỊ THIẾU DO LỖI OCR — KHÔNG ĐỌC ĐƯỢC]", italic=True)

    add_article_title(doc, "ĐIỀU 6: GIẢI QUYẾT TRANH CHẤP")
    add_paragraph_text(doc, "Tranh chấp được giải quyết qua thương lượng trong 15 ngày.")
    add_paragraph_text(doc, "Nếu không đạt, đưa ra Tòa án nhân dân TP. HCM.")

    add_signature_block(doc, "Ngày ký: 10 tháng 02 năm 2026")

    return doc


def build_contract_003():
    """Hợp đồng 003: vận hành mạng, 3 cờ đỏ."""
    doc = setup_document()

    add_title(doc, "HỢP ĐỒNG DỊCH VỤ VẬN HÀNH\nVÀ BẢO TRÌ HỆ THỐNG MẠNG")
    add_contract_number(doc, "Số: HD-VH-2026-003")
    add_disclaimer(doc)

    add_party_info(doc, "Bên A:", [
        "Công ty Cổ phần Viễn thông Mô phỏng (VTN-Sim)",
        "Địa chỉ: 123 đường Nguyễn Văn Cừ, Quận 5, TP. HCM",
        "Mã số thuế: 0301234567-001 (mô phỏng)",
        "Người đại diện: Hoàng Văn Em — Giám đốc",
    ])
    add_paragraph_text(doc, "")
    add_party_info(doc, "Bên B:", [
        "Công ty TNHH Dịch vụ Kỹ thuật Demo",
        "Địa chỉ: 321 đường Lý Thường Kiệt, Quận 10, TP. HCM",
        "Mã số thuế: 0303456789-001 (mô phỏng)",
        "Người đại diện: Võ Thị Phúc — Trưởng phòng Kinh doanh",
    ])

    add_paragraph_text(doc, "")

    add_article_title(doc, "ĐIỀU 1: ĐỐI TƯỢNG HỢP ĐỒNG")
    add_paragraph_text(doc, "Bên B cung cấp dịch vụ vận hành và bảo trì hệ thống mạng cho Bên A, bao gồm:")
    add_bullet(doc, "Giám sát hệ thống 24/7/365")
    add_bullet(doc, "Bảo trì định kỳ hàng tháng")
    add_bullet(doc, "Xử lý sự cố trong vòng 2 giờ kể từ khi nhận thông báo")
    add_bullet(doc, "Báo cáo hoạt động hàng tháng")

    add_article_title(doc, "ĐIỀU 2: THỜI HẠN HỢP ĐỒNG")
    add_paragraph_text(doc, "Hợp đồng có hiệu lực kể từ ngày 01 tháng 03 năm 2026.")
    add_paragraph_text(doc, "Hợp đồng có thời hạn 12 tháng, hết hạn ngày 28 tháng 02 năm 2027.")
    add_paragraph_text(doc, "Hợp đồng sẽ tự động gia hạn thêm 12 tháng mà không cần thông báo trước trừ khi một bên thông báo bằng văn bản chấm dứt trước 15 ngày so với ngày hết hạn.")

    add_article_title(doc, "ĐIỀU 3: GIÁ TRỊ VÀ THANH TOÁN")
    add_paragraph_text(doc, "Giá trị hợp đồng: 1.800.000.000 VNĐ (Một tỷ tám trăm triệu đồng).")
    add_paragraph_text(doc, "Thanh toán hàng tháng: 150.000.000 VNĐ.")
    add_paragraph_text(doc, "Thanh toán trong vòng 10 ngày kể từ ngày nhận hóa đơn.")

    add_article_title(doc, "ĐIỀU 4: ĐIỀU KHOẢN PHẠT VI PHẠM")
    add_paragraph_text(doc, "Nếu Bên B xử lý sự cố vượt quá 2 giờ:")
    add_bullet(doc, "Phạt 5% giá trị hợp đồng hàng tháng cho mỗi giờ vượt.")
    add_bullet(doc, "Không giới hạn mức phạt tối đa.")
    add_paragraph_text(doc, "Nếu Bên A chậm thanh toán:")
    add_bullet(doc, "Phạt 0.1% giá trị chưa thanh toán cho mỗi ngày chậm.")
    add_bullet(doc, "Sau 60 ngày chậm, Bên B có quyền đơn phương chấm dứt hợp đồng.")

    add_article_title(doc, "ĐIỀU 5: GIỚI HẠN TRÁCH NHIỆM")
    add_paragraph_text(doc, "Tổng trách nhiệm bồi thường của Bên B trong mọi trường hợp không vượt quá giá trị hợp đồng của 01 (một) tháng.")
    add_paragraph_text(doc, "Bên B không chịu trách nhiệm đối với thiệt hại gián tiếp, mất mát doanh thu hoặc chi phí cơ hội.")

    add_article_title(doc, "ĐIỀU 6: CHẤM DỨT HỢP ĐỒNG")
    add_paragraph_text(doc, "Hợp đồng chấm dứt khi:")
    add_bullet(doc, "Hết hạn theo Điều 2 (xem điều khoản tự gia hạn).")
    add_bullet(doc, "Thỏa thuận của cả hai bên.")
    add_bullet(doc, "Bên A vi phạm thanh toán quá 60 ngày liên tục.")

    add_article_title(doc, "ĐIỀU 7: ĐIỀU KHOẢN BẢO MẬT")
    add_paragraph_text(doc, "Bên B cam kết bảo mật thông tin hệ thống mạng và dữ liệu vận hành của Bên A.")
    add_paragraph_text(doc, "Nghĩa vụ bảo mật áp dụng trong thời hạn hợp đồng và 1 năm sau khi chấm dứt.")

    add_article_title(doc, "ĐIỀU 8: BẢO HIỂM")
    add_paragraph_text(doc, "Bên B phải mua bảo hiểm trách nhiệm nghề nghiệp với mức bảo hiểm tối thiểu bằng 50% giá trị hợp đồng.")

    add_signature_block(doc, "Ngày ký: 20 tháng 02 năm 2026")

    return doc


def build_contract_004():
    """Hợp đồng 004: thuê kênh quốc tế, SLA 99.99%."""
    doc = setup_document()

    add_title(doc, "HỢP ĐỒNG THUÊ KÊNH\nTRUYỀN DẪN QUỐC TẾ")
    add_contract_number(doc, "Số: HD-TD-2026-004")
    add_disclaimer(doc)

    add_party_info(doc, "Bên A:", [
        "Công ty Cổ phần Viễn thông Mô phỏng (VTN-Sim)",
        "Địa chỉ: 123 đường Nguyễn Văn Cừ, Quận 5, TP. HCM",
        "Mã số thuế: 0301234567-001 (mô phỏng)",
        "Người đại diện: Ngô Văn Hùng — Giám đốc Khối Doanh nghiệp",
    ])
    add_paragraph_text(doc, "")
    add_party_info(doc, "Bên B:", [
        "Công ty TNHH Truyền dẫn Quốc tế Demo",
        "Địa chỉ: 567 đường Điện Biên Phủ, Quận Bình Thạnh, TP. HCM",
        "Mã số thuế: 0307890123-001 (mô phỏng)",
        "Người đại diện: Lý Thị Khánh — Giám đốc Kinh doanh",
    ])

    add_paragraph_text(doc, "")

    add_article_title(doc, "ĐIỀU 1: ĐỐI TƯỢNG HỢP ĐỒNG")
    add_paragraph_text(doc, "Bên B cung cấp cho Bên A dịch vụ thuê kênh truyền dẫn quốc tế với thông số kỹ thuật:")
    add_bullet(doc, "Băng thông: 10 Gbps STM-64")
    add_bullet(doc, "Tuyến: Hồ Chí Minh — Singapore (qua cáp quang biển APG)")
    add_bullet(doc, "Cổng kết nối: 2 cổng 10GE tại POP Tân Thuận và POP Singapore")
    add_bullet(doc, "Độ trễ một chiều: ≤ 15 ms")
    add_bullet(doc, "Tỷ lệ mất gói tin: ≤ 0.01%")

    add_article_title(doc, "ĐIỀU 2: CAM KẾT MỨC DỊCH VỤ (SLA)")
    add_paragraph_text(doc, "Bên B cam kết mức dịch vụ tối thiểu:")
    add_paragraph_text(doc, "a) Thời gian sẵn sàng — Uptime SLA:", indent=0.5)
    add_bullet(doc, "Mục tiêu: 99.99% (tối đa 52.56 phút ngừng dịch vụ/năm)")
    add_bullet(doc, "Thời gian bảo trì định kỳ: tối đa 4 giờ/tháng, lịch trước 72 giờ")
    add_bullet(doc, "Thời gian bảo trì không tính vào SLA downtime")
    add_paragraph_text(doc, "b) Thời gian khôi phục sự cố — Disaster Recovery:", indent=0.5)
    add_bullet(doc, "RTO (Recovery Time Objective): < 4 giờ kể từ khi phát hiện sự cố")
    add_bullet(doc, "RPO (Recovery Point Objective): < 1 giờ")
    add_bullet(doc, "Kế hoạch chuyển đổi dự phòng: dự phòng 2 tuyến cáp quang (APG + AAG)")
    add_paragraph_text(doc, "c) Báo cáo hiệu suất:", indent=0.5)
    add_bullet(doc, "Báo cáo hàng tháng: uptime thực tế, số lần sự cố, thời gian khôi phục, độ trễ trung bình")
    add_bullet(doc, "Báo cáo quý: xu hướng hiệu suất, khuyến nghị cải thiện")
    add_bullet(doc, "Báo cáo sự cố riêng: trong 48 giờ sau mỗi sự cố, kèm phân tích nguyên nhân gốc")

    add_article_title(doc, "ĐIỀU 3: THỜI HẠN HỢP ĐỒNG")
    add_paragraph_text(doc, "Hợp đồng có hiệu lực kể từ ngày 01 tháng 06 năm 2026.")
    add_paragraph_text(doc, "Hợp đồng có thời hạn 24 tháng, hết hạn vào ngày 31 tháng 05 năm 2028.")
    add_paragraph_text(doc, "Tự động gia hạn thêm 12 tháng trừ khi một bên thông báo bằng văn bản trước 90 ngày.")

    add_article_title(doc, "ĐIỀU 4: GIÁ TRỊ VÀ THANH TOÁN")
    add_paragraph_text(doc, "Giá trị hợp đồng: 4.800.000.000 VNĐ/năm (Bốn tỷ tám trăm triệu đồng/năm).")
    add_paragraph_text(doc, "Thanh toán hàng tháng: 400.000.000 VNĐ.")
    add_paragraph_text(doc, "Thanh toán trong vòng 15 ngày kể từ ngày nhận hóa đơn.")

    add_article_title(doc, "ĐIỀU 5: ĐIỀU KHOẢN PHẠT")
    add_paragraph_text(doc, "Nếu Bên B không đạt SLA uptime 99.99% trong một tháng:")
    add_bullet(doc, "Phạt 0.1% giá trị hợp đồng hàng tháng cho mỗi phút ngừng dịch vụ vượt quá mức cho phép.")
    add_bullet(doc, "Mức cho phép: 4.38 phút/tháng (= 52.56 phút/năm ÷ 12).")
    add_bullet(doc, "Tối đa phạt không quá 25% giá trị hợp đồng hàng tháng.")
    add_paragraph_text(doc, "Nếu Bên B không khôi phục sự cố trong RTO 4 giờ:")
    add_bullet(doc, "Phạt thêm 50.000.000 VNĐ cho mỗi giờ vượt RTO.")
    add_bullet(doc, "Không giới hạn mức phạt RTO.")
    add_paragraph_text(doc, "Nếu Bên A chậm thanh toán:")
    add_bullet(doc, "Phạt 0.05% giá trị chưa thanh toán cho mỗi ngày chậm.")
    add_bullet(doc, "Sau 30 ngày chậm, Bên B có quyền tạm ngưng dịch vụ sau khi thông báo trước 5 ngày.")

    add_article_title(doc, "ĐIỀU 6: CHẤM DỨT HỢP ĐỒNG")
    add_paragraph_text(doc, "Hợp đồng chấm dứt khi:")
    add_bullet(doc, "Hết hạn theo Điều 3 mà không gia hạn.")
    add_bullet(doc, "Thỏa thuận bằng văn bản của cả hai bên.")
    add_bullet(doc, "Một bên vi phạm nghiêm trọng và không khắc phục trong 30 ngày.")
    add_bullet(doc, "SLA vi phạm 3 tháng liên tiếp: Bên A có quyền chấm dứt ngay lập tức mà không bồi thường.")

    add_article_title(doc, "ĐIỀU 7: BẢO MẬT THÔNG TIN")
    add_paragraph_text(doc, "Cả hai bên cam kết bảo mật:")
    add_bullet(doc, "Dữ liệu lưu lượng truyền dẫn, thông tin khách hàng, cấu hình kỹ thuật.")
    add_bullet(doc, "Báo cáo hiệu suất và SLA không được chia sẻ cho bên thứ ba mà không có đồng ý bằng văn bản.")
    add_bullet(doc, "Nghĩa vụ bảo mật kéo dài 3 năm sau khi hợp đồng hết hạn.")

    add_article_title(doc, "ĐIỀU 8: BẢO HIỂM VÀ BỒI THƯỜNG")
    add_paragraph_text(doc, "Bên B phải duy trì bảo hiểm trách nhiệm nghề nghiệp tối thiểu 1.000.000.000 VNĐ.")
    add_paragraph_text(doc, "Tổng trách nhiệm bồi thường của Bên B không vượt quá giá trị hợp đồng của 03 tháng.")

    add_article_title(doc, "ĐIỀU 9: GIẢI QUYẾT TRANH CHẤP")
    add_paragraph_text(doc, "Mọi tranh chấp phát sinh sẽ được giải quyết qua thương lượng trong 30 ngày.")
    add_paragraph_text(doc, "Nếu không đạt, tranh chấp sẽ được đưa ra Trọng tài thương mại TP. HCM.")
    add_paragraph_text(doc, "Mỗi bên chịu chi phí trọng tài của mình, trừ trường hợp trọng tài quyết định khác.")

    add_signature_block(doc, "Ngày ký: 15 tháng 05 năm 2026")

    return doc


def main():
    output_dir = os.path.dirname(os.path.abspath(__file__))
    contracts_dir = os.path.join(output_dir, "contracts")
    os.makedirs(contracts_dir, exist_ok=True)

    builders = [
        ("contract-001.docx", build_contract_001),
        ("contract-002.docx", build_contract_002),
        ("contract-003-risky.docx", build_contract_003),
        ("contract-004-telecom-sla.docx", build_contract_004),
    ]

    for filename, builder in builders:
        doc = builder()
        filepath = os.path.join(contracts_dir, filename)
        doc.save(filepath)
        print(f"OK: {filepath}")

    print(f"\nHoàn thành: {len(builders)} file DOCX")


if __name__ == "__main__":
    main()
